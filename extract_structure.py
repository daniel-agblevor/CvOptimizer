import json
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxBlueprintExtractor:
    def __init__(self, file_path):
        self.file_path = file_path
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        self.doc = Document(file_path)
        self.blueprint = {
            "metadata": {},
            "sections": [],
            "content_structure": []
        }

    def _get_color_hex(self, color_obj):
        """Converts docx color object to Hex string."""
        if color_obj and color_obj.rgb:
            return str(color_obj.rgb)
        return "Auto/Inherited"

    def _parse_run(self, run):
        """Extracts formatting details from a specific text run."""
        return {
            "text": run.text,
            "style": {
                "font": run.font.name,
                "size": run.font.size.pt if run.font.size else "Inherited",
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "strike": run.font.strike,
                "color": self._get_color_hex(run.font.color)
            }
        }

    def _parse_paragraph(self, para, index):
        """Parses a paragraph and its children runs."""
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "Left",
            WD_ALIGN_PARAGRAPH.CENTER: "Center",
            WD_ALIGN_PARAGRAPH.RIGHT: "Right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "Justify"
        }
        
        # Determine alignment (defaulting to Left if None)
        align = alignment_map.get(para.alignment, "Left/Inherited")

        structure = {
            "type": "paragraph",
            "index": index,
            "style_name": para.style.name,
            "alignment": align,
            "runs": [self._parse_run(run) for run in para.runs if run.text.strip()]
        }
        return structure

    def _parse_table(self, table, index):
        """Parses table structure."""
        table_data = {
            "type": "table",
            "index": index,
            "rows": []
        }
        
        for i, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                # Recursively parse paragraphs inside table cells
                cell_content = [self._parse_paragraph(p, 0) for p in cell.paragraphs if p.text.strip()]
                row_data.append(cell_content)
            table_data["rows"].append(row_data)
            
        return table_data

    def extract(self):
        # 1. Extract Metadata
        core_props = self.doc.core_properties
        self.blueprint["metadata"] = {
            "author": core_props.author,
            "created": str(core_props.created),
            "modified": str(core_props.modified),
            "last_modified_by": core_props.last_modified_by,
            "revision": core_props.revision
        }

        # 2. Extract Section Layouts (Page setup)
        for i, section in enumerate(self.doc.sections):
            self.blueprint["sections"].append({
                "section_index": i,
                "orientation": section.orientation,
                "page_width": section.page_width.pt,
                "page_height": section.page_height.pt,
                "margin_left": section.left_margin.pt,
                "margin_right": section.right_margin.pt
            })

        # 3. Extract Content (Iterating in document order)
        # Note: python-docx separates tables and paragraphs. 
        # For a pure sequential blueprint, we iterate the element body (advanced).
        # Here we extract paragraphs first, then tables for simplicity, 
        # or we can inspect the child elements of the document body.
        
        # Using a simpler sequential logic by iterating over valid children
        for element in self.doc.element.body:
            if element.tag.endswith('p'): # Paragraph
                # Find the paragraph object in doc.paragraphs that matches this element
                # This is a linear search; efficient enough for blueprinting
                para_obj = next((p for p in self.doc.paragraphs if p._element is element), None)
                if para_obj and para_obj.text.strip():
                    self.blueprint["content_structure"].append(
                        self._parse_paragraph(para_obj, len(self.blueprint["content_structure"]))
                    )
            elif element.tag.endswith('tbl'): # Table
                table_obj = next((t for t in self.doc.tables if t._element is element), None)
                if table_obj:
                    self.blueprint["content_structure"].append(
                        self._parse_table(table_obj, len(self.blueprint["content_structure"]))
                    )

        return self.blueprint

    def save_json(self, output_path):
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.blueprint, f, indent=4)
        print(f"Blueprint saved to: {output_path}")

# --- Usage Example ---
if __name__ == "__main__":
    # create a dummy file for testing if one doesn't exist
    if not os.path.exists("test_doc.docx"):
        doc = Document()
        doc.add_heading('Document Blueprint Test', 0)
        p = doc.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True
        doc.save("test_doc.docx")

    try:
        extractor = DocxBlueprintExtractor("test_doc.docx")
        data = extractor.extract()
        extractor.save_json("blueprint.json")
    except Exception as e:
        print(f"Error: {e}")