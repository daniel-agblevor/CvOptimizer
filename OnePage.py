import os
import json
import re
import logging
from decouple import config
import google.generativeai as genai
from docx import Document
from docx.text.paragraph import Paragraph

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class OnePageCVBuilder:
    # Fields that typically require bullet point formatting
    BULLET_POINT_SUFFIXES = ["_DESC", "_ACHIEVEMENTS", "_RESPONSIBILITIES", "_DETAILS"]

    def __init__(self, api_key, template_path, output_path="Final_OnePage_CV.docx"):
        logging.info("Initializing OnePageCVBuilder...")
        self.template_path = template_path
        self.output_path = output_path
        
        # Configure Gemini
        logging.info("Configuring Gemini API...")
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.5-flash')
        
        # 1. Analyze Template to Build Schema
        logging.info(f"Analyzing template: {template_path}")
        self.all_blocks = self._get_all_blocks(template_path)
        self.static_placeholders, self.dynamic_sections = self._parse_template_structure()
        self.schema_guidelines = self._build_dynamic_schema()
        logging.info("Template analysis complete.")

    def _get_all_blocks(self, file_path):
        """Helper to get all paragraphs, including those in tables."""
        doc = Document(file_path)
        all_blocks = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_blocks.extend(cell.paragraphs)
        return all_blocks
    
    def _parse_template_structure(self):
        """
        Scans the template for:
        1. Static fields: {{NAME}}, {{SUMMARY}}
        2. Dynamic repeatable sections: {{JOB1_TITLE}}, {{CERT2_NAME}}
        
        Returns:
            static_set: Set of static keys.
            dynamic_map: Dictionary defining repeatable sections.
                e.g., {'JOB': {'max_index': 4, 'fields': {'TITLE', 'COMPANY'}}}
        """
        static_set = set()
        dynamic_map = {} # {PREFIX: {max_index: int, fields: set}}
        
        for block in self.all_blocks:
            # Find all potential keys inside {{...}}
            # We look for simple {{KEY}} or indexed {{PREFIX#_KEY}}
            matches = re.findall(r"\{\{([A-Z0-9_]+)\}\}", block.text)
            
            for full_key in matches:
                # Regex to detect indexed patterns like JOB1_TITLE or CERT2_DATE
                # Group 1: Prefix (JOB), Group 2: Index (1), Group 3: Field (TITLE)
                indexed_match = re.match(r"^([A-Z]+)(\d+)_([A-Z_]+)$", full_key)
                
                if indexed_match:
                    prefix, index, field = indexed_match.groups()
                    index = int(index)
                    
                    if prefix not in dynamic_map:
                        dynamic_map[prefix] = {'max_index': 0, 'fields': set()}
                    
                    # Update max index found (to know how many slots the template has)
                    dynamic_map[prefix]['max_index'] = max(dynamic_map[prefix]['max_index'], index)
                    dynamic_map[prefix]['fields'].add(field)
                else:
                    # It's a static field like NAME, EMAIL, SKILLS
                    static_set.add(full_key)
        
        return list(static_set), dynamic_map

    def _build_dynamic_schema(self):
        """Creates a tailored JSON schema for the AI based on the specific template structure."""
        schema = {}
        
        # 1. Define Static Fields
        for key in self.static_placeholders:
            schema[key] = {"type": "string", "description": f"Content for {key}. If not found, return empty string."}
        
        # 2. Define Dynamic Array Fields
        for prefix, meta in self.dynamic_sections.items():
            field_structure = {field: {"type": "string"} for field in meta['fields']}
            
            description = f"List of dictionaries for {prefix} (e.g., Work Experience, Certificates). " \
                          f"Extract the most relevant items."
            
            # Special instruction for commonly known sections
            if prefix == "JOB" or prefix == "EXP":
                description += " Limit to 3-4 most relevant roles. Use active verbs."
            
            schema[prefix] = {
                "type": "array",
                "description": description,
                "items": {
                    "type": "object",
                    "properties": field_structure
                }
            }
        
        return schema

    def extract_text_from_candidate(self, file_path):
        """Simple extractor for DOCX."""
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)

    def optimize_content(self, raw_text):
        """Uses Gemini to generate structured data matching the template."""
        logging.info("Extracting and structuring data with Gemini...")
        
        # Keys the AI must generate
        expected_keys = self.static_placeholders + list(self.dynamic_sections.keys())

        prompt = f"""
        You are an expert CV parser and writer.
        
        TASK:
        1. Parse the RAW CANDIDATE CV text below.
        2. Map the content to the required JSON schema structure defined by the template placeholders.
        3. For Array fields (like JOB, CERT), return a list of objects.
        4. Use concise, professional language.
        5. For description fields (ending in _DESC or _ACHIEVEMENTS), use 3-5 bullet points separated by newlines (\\n).
        6. If a field is missing in the CV, return an empty string "" or empty list [].
        
        RAW CANDIDATE TEXT:
        {raw_text[:10000]}
        
        REQUIRED SCHEMA (Target Structure):
        {json.dumps(self.schema_guidelines, indent=2)}
        
        OUTPUT FORMAT:
        Return a single valid JSON object with keys: {expected_keys}
        """
        
        response = self.model.generate_content(prompt)
        try:
            json_str = response.text.replace('```json', '').replace('```', '').strip()
            logging.info(f"AI Output Preview: {json_str[:200]}...")
            logging.info("Successfully extracted structured data from Gemini.")
            return json.loads(json_str)
        except json.JSONDecodeError:
            logging.error(f"AI did not return valid JSON. Full response: {response.text}")
            return {}

    def _replace_text_preserve_style(self, element, placeholder, new_text):
        """Replaces placeholder in a paragraph or run while preserving style."""
        if isinstance(element, Paragraph) and placeholder in element.text:
            # Attempt to replace within runs to keep bold/italic/color
            for run in element.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, new_text)
                    return
            # Fallback
            element.text = element.text.replace(placeholder, new_text)

    def _format_bullet_points(self, text):
        """Helper to format text as visual bullet points."""
        if not text:
            return ""
        # If it already looks like a list, strictly format it
        formatted = text.replace('\n', '\n• ')
        if not formatted.startswith('• '):
            formatted = '• ' + formatted
        return formatted

    def generate_document(self, candidate_path):
        # 1. Ingest
        raw_text = self.extract_text_from_candidate(candidate_path)
        
        # 2. Structure
        data_map = self.optimize_content(raw_text)
        
        # 3. Load Template
        logging.info(f"--- Loading template: {self.template_path} ---")
        doc = Document(self.template_path)
        
        # Reload blocks from the fresh doc object for modification
        all_blocks_to_modify = self._get_all_blocks(self.template_path)
        
        logging.info("--- Injecting data... ---")

        # --- PHASE 1: STATIC FIELDS ---
        for key in self.static_placeholders:
            placeholder = f"{{{{{key}}}}}"
            value = data_map.get(key, "")
            
            for block in all_blocks_to_modify:
                if placeholder in block.text:
                    if not value and isinstance(block, Paragraph):
                        block.clear() # Cleanup empty static fields
                    else:
                        self._replace_text_preserve_style(block, placeholder, str(value))

        # --- PHASE 2: DYNAMIC REPEATABLE SECTIONS ---
        # Iterate over every detected dynamic section type (JOB, CERT, PROJ, etc.)
        for prefix, meta in self.dynamic_sections.items():
            section_data = data_map.get(prefix, []) # Get list of dicts from AI (e.g., list of jobs)
            max_slots = meta['max_index']
            
            logging.info(f"Processing section '{prefix}': Found {len(section_data)} items for {max_slots} slots.")

            # Iterate through the template slots (1 to max_slots)
            for i in range(1, max_slots + 1):
                item_data = section_data[i-1] if i <= len(section_data) else {}
                
                # For every field known to this section (TITLE, COMPANY, etc.)
                for field in meta['fields']:
                    # Construct the specific placeholder: {{JOB1_TITLE}}
                    placeholder = f"{{{{{prefix}{i}_{field}}}}}"
                    
                    # Extract value from the AI dictionary
                    value = item_data.get(field, "")
                    
                    # Apply bullet formatting if applicable
                    if any(field.endswith(suffix) for suffix in self.BULLET_POINT_SUFFIXES):
                         value = self._format_bullet_points(value)

                    # Scan document to find and replace
                    for block in all_blocks_to_modify:
                        if placeholder in block.text:
                            if not value and isinstance(block, Paragraph):
                                # CLEANUP: If this slot is unused (e.g., Job 3 when only 2 exist), clear it
                                block.clear()
                            else:
                                self._replace_text_preserve_style(block, placeholder, str(value))

        # 4. Save
        doc.save(self.output_path)
        logging.info(f"--- Document generated: {self.output_path} ---")

if __name__ == "__main__":
    # CONFIGURATION
    API_KEY = config("GOOGLE_API_KEY") 
    CANDIDATE_FILE = "CandidateCV.docx"
    TEMPLATE_FILE = "Standard.docx"
    
    if os.path.exists(CANDIDATE_FILE) and os.path.exists(TEMPLATE_FILE):
        logging.info("Starting CV generation process...")
        builder = OnePageCVBuilder(API_KEY, TEMPLATE_FILE)
        builder.generate_document(CANDIDATE_FILE)
        logging.info("CV generation process finished.")
    else:
        logging.error("Please ensure candidate and template files exist.")