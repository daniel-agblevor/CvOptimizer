import os
import logging
from decouple import config
import google.generativeai as genai
from docx import Document
from docx.shared import Pt

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class CvOptimizer:
    def __init__(self, api_key, jd_path, cv_path, output_path="OptimizedCV.docx"):
        logging.info("Initializing CvOptimizer...")
        self.jd_path = jd_path
        self.cv_path = cv_path
        self.output_path = output_path
        
        # Configure Gemini
        logging.info("Configuring Gemini API...")
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.5-pro')
        
        # Load Job Description
        logging.info(f"Loading job description from: {self.jd_path}")
        with open(self.jd_path, 'r', encoding='utf-8') as f:
            self.job_description = f.read()

        # Load CV Document Object
        if not os.path.exists(self.cv_path):
            logging.error(f"CV file not found: {self.cv_path}")
            raise FileNotFoundError(f"CV file not found: {self.cv_path}")
        logging.info(f"Loading CV from: {self.cv_path}")
        self.doc = Document(self.cv_path)

    def _is_meaningful_content(self, text):
        """
        Filters out short headers, dates, or empty lines to save API tokens
        and prevent the AI from hallucinating on 'Name' or 'Address'.
        """
        if not text:
            return False
        # If it's very short (likely a name or date), skip optimization
        if len(text.split()) < 4:
            return False
        return True

    def _optimize_text_segment(self, original_text):
        """
        Sends a specific paragraph/bullet to Gemini for rewriting.
        """
        prompt = f"""
        CONTEXT: You are a professional CV writer. 
        TASK: Rewrite the following CV text segment to align better with the JOB DESCRIPTION provided below.
        
        GUIDELINES:
        1. Keep the length and tone similar to the original text.
        2. Use keywords from the Job Description where natural.
        3. Do not invent facts. Only rephrase existing experience.
        4. Return ONLY the rewritten text. No markdown, no quotes, no explanations.
        
        JOB DESCRIPTION:
        {self.job_description[:2000]}... (truncated for context)

        ORIGINAL CV FRAGMENT:
        "{original_text}"
        """
        
        try:
            response = self.model.generate_content(prompt)
            cleaned_response = response.text.strip().replace('"', '')
            logging.info(f"Optimized Text --- Original: '{original_text}' -> Optimized: '{cleaned_response}'")
            return cleaned_response
        except Exception as e:
            logging.error(f"API Error on segment: {e}")
            return original_text # Fallback to original if API fails

    def _replace_paragraph_text(self, paragraph, new_text):
        """
        Replaces text in a paragraph while attempting to preserve the 
        formatting of the original first run (font, size, bolding).
        """
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return

        # Capture style from the first run of the original paragraph
        ref_run = paragraph.runs[0]
        p_style = paragraph.style
        
        # Clear existing content
        paragraph.clear()
        
        # Add new text and re-apply basic formatting from reference
        new_run = paragraph.add_run(new_text)
        new_run.font.size = Pt(11)
        new_run.bold = ref_run.bold
        new_run.italic = ref_run.italic
        new_run.underline = ref_run.underline
        
        # Note: Complex mixed styles (e.g., one bold word in a sentence) 
        # are flattened here to ensure the new sentence structure flows correctly.

    def process(self):
        logging.info("--- Starting CV Optimization ---")
        
        # 1. Traverse Document Body (Standard Paragraphs)
        total_paragraphs = len(self.doc.paragraphs)
        logging.info(f"Processing {total_paragraphs} body paragraphs...")
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if self._is_meaningful_content(text):
                logging.info(f"Optimizing paragraph {i+1}...")
                optimized_text = self._optimize_text_segment(text)
                self._replace_paragraph_text(para, optimized_text)

        # 2. Traverse Tables (Many CVs use tables for layout)
        logging.info("Processing tables...")
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if self._is_meaningful_content(text):
                            logging.info("Optimizing text in table cell...")
                            optimized_text = self._optimize_text_segment(text)
                            self._replace_paragraph_text(para, optimized_text)

        # 3. Save the result
        self.doc.save(self.output_path)
        logging.info(f"--- Success! Optimized CV saved to {self.output_path} ---")

# --- Execution ---
if __name__ == "__main__":
    # Create dummy files for demonstration if they don't exist
    if not os.path.exists("job_description.txt"):
        with open("job_description.txt", "w") as f:
            f.write("Looking for a Python Developer with experience in AI, automation, and data processing.")
    
    # You would provide your actual file paths here
    API_KEY = config("GOOGLE_API_KEY") # Replace this
    
    try:
        optimizer = CvOptimizer(
            api_key=API_KEY,
            jd_path="job_description.txt",
            cv_path="CandidateCV.docx", # Ensure this file exists
            output_path="OptimizedCV.docx"
        )
        optimizer.process()
    except Exception as e:
        logging.error(f"An error occurred during the optimization process: {e}")